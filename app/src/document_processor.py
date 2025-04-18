import re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
import importlib.util
import os

# Kiểm tra và tải spaCy
try:
    import spacy
    try:
        # Kiểm tra xem mô hình đã được cài đặt chưa
        spacy_model_name = 'en_core_web_sm'
        if importlib.util.find_spec(spacy_model_name) is not None:
            # Mô hình đã được cài đặt dưới dạng package Python
            nlp = spacy.load(spacy_model_name)
            print(f"Đã tải thành công mô hình {spacy_model_name}")
        else:
            # Thử tải từ đường dẫn khác
            try:
                nlp = spacy.load(spacy_model_name)
                print(f"Đã tải thành công mô hình {spacy_model_name}")
            except OSError:
                print(f"Không tìm thấy mô hình {spacy_model_name}. Đang tải mô hình nhỏ mặc định...")
                # Sử dụng mô hình nhỏ mặc định nếu không tìm thấy mô hình chính
                nlp = spacy.blank('en')
                print("Đã tải mô hình nhỏ (blank model)")
    except Exception as e:
        print(f"Lỗi khi tải mô hình spaCy: {str(e)}")
        # Tạo mô hình trống nếu không thể tải được mô hình
        nlp = spacy.blank('en')
        print("Đã tải mô hình trống (blank model)")
except ImportError as e:
    print(f"Lỗi: Không thể import thư viện spaCy. {str(e)}")
    print("Vui lòng cài đặt spaCy bằng lệnh 'pip install spacy' và chạy 'python -m spacy download en_core_web_sm'")
    raise ImportError("Không thể sử dụng ứng dụng khi thiếu thư viện spaCy")


def allowed_file(filename, allowed_extensions):
    """Kiểm tra xem tệp có định dạng được phép không."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def extract_text_from_doc(file_path):
    """Trích xuất văn bản từ tài liệu Word hoặc tệp văn bản."""
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    return None


def analyze_text(content):
    """Phân tích văn bản với spaCy để trích xuất thông tin hữu ích."""
    try:
        doc = nlp(content)
        
        # Kiểm tra xem mô hình có khả năng phân tích văn bản không
        has_nlp_capabilities = hasattr(doc[0], 'pos_') if len(doc) > 0 else False
        
        if has_nlp_capabilities:
            # Phân tích NLP đầy đủ
            analysis = {
                'entities': [{'text': ent.text, 'label': ent.label_} for ent in doc.ents],
                'sentences': len(list(doc.sents)),
                'tokens': len(doc),
                'keywords': [token.text for token in doc if token.pos_ in ("NOUN", "ADJ", "VERB") 
                            and not token.is_stop and len(token.text) > 3][:50],
                'noun_chunks': [chunk.text for chunk in doc.noun_chunks][:50]
            }
        else:
            # Phân tích cơ bản khi không có khả năng NLP đầy đủ
            words = content.split()
            sentences = [s.strip() for s in re.split(r'[.!?]+', content) if s.strip()]
            
            # Ước tính các từ khóa bằng tần suất xuất hiện
            word_freq = {}
            for word in words:
                word = re.sub(r'[^\w\s]', '', word.lower())
                if word and len(word) > 3:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # Lấy top từ khóa theo tần suất
            keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            keywords = [word for word, freq in keywords[:50]]
            
            analysis = {
                'entities': [],
                'sentences': len(sentences),
                'tokens': len(words),
                'keywords': keywords,
                'noun_chunks': []
            }
        
        return analysis
    except Exception as e:
        # Fallback nếu có lỗi: trả về phân tích cơ bản
        print(f"Lỗi trong quá trình phân tích văn bản: {str(e)}")
        words = content.split()
        sentences = [s.strip() for s in re.split(r'[.!?]+', content) if s.strip()]
        
        return {
            'entities': [],
            'sentences': len(sentences),
            'tokens': len(words),
            'keywords': [],
            'noun_chunks': []
        }


def format_text_with_spacy(content, formatting_options):
    """Định dạng nội dung văn bản bằng spaCy """
    try:
        doc = nlp(content)
        
        # Kiểm tra xem mô hình có khả năng phân tích văn bản không
        has_nlp_capabilities = hasattr(doc[0], 'pos_') if len(doc) > 0 else False
        
        # Phân tích cấu trúc văn bản
        paragraphs = []
        headings = []
        
        if has_nlp_capabilities:
            # Xác định các tiêu đề và đoạn văn với NLP
            current_paragraph = []
            for sent in doc.sents:
                sent_text = sent.text.strip()
                
                # Phát hiện tiêu đề dựa trên độ dài, hoa/thường, số từ
                is_heading = (len(sent_text) < 100 and 
                            len(sent_text.split()) < 10 and 
                            any(token.is_title for token in sent if token.is_alpha))
                
                if is_heading:
                    # Nếu có đoạn văn đang mở, đóng lại
                    if current_paragraph:
                        paragraphs.append(' '.join(current_paragraph))
                        current_paragraph = []
                    
                    # Thêm tiêu đề
                    headings.append(sent_text)
                else:
                    current_paragraph.append(sent_text)
            
            # Thêm đoạn văn cuối cùng nếu có
            if current_paragraph:
                paragraphs.append(' '.join(current_paragraph))
        else:
            # Fallback khi không có khả năng NLP: chia theo đoạn và thêm tiêu đề tự động
            raw_paragraphs = content.split('\n\n')
            
            # Tìm các tiêu đề tiềm năng (dòng ngắn)
            for para in raw_paragraphs:
                para = para.strip()
                if para:
                    if len(para) < 100 and len(para.split()) < 10:
                        headings.append(para)
                    else:
                        paragraphs.append(para)
        
        # Đảm bảo có ít nhất một đoạn văn nếu không phát hiện được đoạn nào
        if not paragraphs and not headings:
            # Nếu không phát hiện được tiêu đề hoặc đoạn văn, coi toàn bộ nội dung là một đoạn văn
            paragraphs = [content]
        
        # Tạo tiêu đề mặc định nếu không có tiêu đề nào được phát hiện
        if not headings and paragraphs:
            # Lấy câu đầu tiên từ đoạn đầu tiên làm tiêu đề
            first_paragraph = paragraphs[0]
            sentences = first_paragraph.split('.')
            if sentences and len(sentences[0]) < 100:
                headings = [sentences[0]]
            else:
                # Nếu không có câu phù hợp, tạo tiêu đề mặc định
                headings = ["Document"]
        
        # Xử lý trích dẫn
        formatted_text = ""
        citation_style = formatting_options['citation_style']
        
        # Định dạng tiêu đề theo thứ tự
        for i, heading in enumerate(headings):
            section_number = i + 1
            formatted_text += f"{section_number}. {heading}\n\n"
            
            # Thêm một đoạn văn sau mỗi tiêu đề nếu có
            if i < len(paragraphs):
                # Định dạng trích dẫn (giả định)
                para = paragraphs[i]
                
                # Định dạng trích dẫn theo kiểu đã chọn
                if citation_style == 'apa':
                    # Định dạng kiểu APA (Author, Year)
                    para = re.sub(r'\(([^)]+), (\d{4})\)', r'(\1, \2)', para)
                elif citation_style == 'mla':
                    # Định dạng kiểu MLA (Author page)
                    para = re.sub(r'\(([^)]+) (\d+)\)', r'(\1 \2)', para)
                elif citation_style == 'chicago':
                    # Định dạng kiểu Chicago (chú thích)
                    para = re.sub(r'\[\d+\]', lambda m: f"^{m.group(0)[1:-1]}", para)
                elif citation_style == 'ieee':
                    # Định dạng kiểu IEEE [số]
                    para = re.sub(r'\((\d+)\)', r'[\1]', para)
                
                formatted_text += para + "\n\n"
        
        # Xử lý các đoạn văn còn lại
        for i in range(len(headings), len(paragraphs)):
            formatted_text += paragraphs[i] + "\n\n"
        
        return formatted_text
    
    except Exception as e:
        # Fallback nếu có lỗi khi xử lý: trả về nội dung gốc với định dạng cơ bản
        print(f"Lỗi trong quá trình xử lý văn bản: {str(e)}")
        
        # Định dạng văn bản đơn giản
        paragraphs = content.split('\n\n')
        formatted_text = "1. Document\n\n"
        
        for para in paragraphs:
            if para.strip():
                formatted_text += para.strip() + "\n\n"
                
        return formatted_text


def create_formatted_document(content, formatting_options, output_path):
    """Tạo tài liệu Word được định dạng dựa trên các tùy chọn định dạng."""
    print(f"Đang áp dụng các tùy chọn định dạng: {formatting_options}")
    
    doc = docx.Document()
    
    # Set document-wide font
    style = doc.styles['Normal']
    font = style.font
    font.name = formatting_options['font_family']
    font.size = Pt(int(formatting_options['font_size']))
    
    # Set document-wide paragraph spacing
    paragraph_format = style.paragraph_format
    if formatting_options['line_spacing'] == '2.0':
        paragraph_format.line_spacing = 2.0
    elif formatting_options['line_spacing'] == '1.5':
        paragraph_format.line_spacing = 1.5
    else:
        paragraph_format.line_spacing = 1.0
    
    # Set margins
    sections = doc.sections
    for section in sections:
        if formatting_options['margin'] == 'normal':
            section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(1)
        elif formatting_options['margin'] == 'wide':
            section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(1.5)
        else:  # narrow
            section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.5)

    # Thêm trang tiêu đề nếu được yêu cầu
    if formatting_options['title_page']:
        # Phân tích văn bản để trích xuất tiêu đề
        doc_nlp = nlp(content[:500])  # Chỉ phân tích 500 ký tự đầu tiên để tìm tiêu đề
        potential_title = next((sent.text for sent in doc_nlp.sents 
                              if len(sent.text) < 100 and sent.text.strip()), "Document Title")
        
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = 1  # Center
        title_run = title_paragraph.add_run(potential_title)
        title_run.font.size = Pt(16)
        title_run.font.name = formatting_options['font_family']
        title_run.font.bold = True
        doc.add_paragraph()

    # Thêm bảng nội dung nếu được yêu cầu
    if formatting_options['table_of_contents']:
        toc_paragraph = doc.add_paragraph()
        toc_paragraph.alignment = 1  # Center
        toc_run = toc_paragraph.add_run("Table of Contents")
        toc_run.font.size = Pt(14)
        toc_run.font.name = formatting_options['font_family']
        toc_run.font.bold = True
        doc.add_paragraph()
        
        # Tạo danh sách mục lục bằng cách phân tích văn bản với spaCy
        doc_nlp = nlp(content)
        headings = []
        
        # Tìm các tiêu đề tiềm năng
        for sent in doc_nlp.sents:
            sent_text = sent.text.strip()
            # Phát hiện tiêu đề dựa trên độ dài, hoa/thường, số từ
            is_heading = (len(sent_text) < 100 and 
                          len(sent_text.split()) < 10 and 
                          any(token.is_title for token in sent if token.is_alpha))
            
            if is_heading:
                headings.append(sent_text)
        
        # Thêm vào mục lục
        for i, heading in enumerate(headings[:10]):  # Giới hạn 10 mục
            toc_entry = doc.add_paragraph()
            toc_entry.add_run(f"{i+1}. {heading}")
            
        doc.add_paragraph()

    try:
        # Phân tích và định dạng nội dung với spaCy
        formatted_text = format_text_with_spacy(content, formatting_options)
        
        # Chia văn bản thành các phần dựa trên tiêu đề
        sections = re.split(r'\n(?=\d+\.)', formatted_text)
        
        # Áp dụng định dạng cho từng phần
        for section in sections:
            if not section.strip():
                continue
                
            # Tách tiêu đề (nếu có) và nội dung
            section_lines = section.strip().split('\n\n', 1)
            
            # Nếu có tiêu đề (dòng đầu tiên chứa số.), tạo đoạn riêng cho tiêu đề
            if section_lines and re.match(r'^\d+\.', section_lines[0]):
                # Thêm tiêu đề
                heading_para = doc.add_paragraph()
                heading_para.style = 'Heading 1'
                heading_run = heading_para.add_run(section_lines[0])
                heading_run.font.size = Pt(14)
                heading_run.font.name = formatting_options['font_family']
                heading_run.font.bold = True
                
                # Thêm nội dung (nếu có)
                if len(section_lines) > 1:
                    content_text = section_lines[1]
                    para = doc.add_paragraph()
                    para.style = 'Normal'
                    run = para.add_run(content_text)
                    run.font.name = formatting_options['font_family']
                    run.font.size = Pt(int(formatting_options['font_size']))
                    
                    # Set line spacing
                    if formatting_options['line_spacing'] == '2.0':
                        para.paragraph_format.line_spacing = 2.0
                    elif formatting_options['line_spacing'] == '1.5':
                        para.paragraph_format.line_spacing = 1.5
                    else:
                        para.paragraph_format.line_spacing = 1.0
            else:
                # Nếu không có tiêu đề, thêm toàn bộ phần vào một đoạn
                para = doc.add_paragraph()
                para.style = 'Normal'
                run = para.add_run(section.strip())
                run.font.name = formatting_options['font_family']
                run.font.size = Pt(int(formatting_options['font_size']))
                
                # Set line spacing
                if formatting_options['line_spacing'] == '2.0':
                    para.paragraph_format.line_spacing = 2.0
                elif formatting_options['line_spacing'] == '1.5':
                    para.paragraph_format.line_spacing = 1.5
                else:
                    para.paragraph_format.line_spacing = 1.0

        # Thêm số trang nếu được yêu cầu
        if formatting_options['page_numbers']:
            for section in doc.sections:
                footer = section.footer
                paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                paragraph.alignment = 1  # Center
                paragraph.text = "Page "
                run = paragraph.add_run()
                run.font.name = formatting_options['font_family']
                paragraph.text = "- Page # -"

        # Thêm tài liệu tham khảo nếu được yêu cầu
        if formatting_options['bibliography']:
            doc.add_paragraph()
            biblio_heading = doc.add_paragraph()
            biblio_heading.alignment = 1  # Center
            biblio_heading_run = biblio_heading.add_run("References")
            biblio_heading_run.font.size = Pt(14)
            biblio_heading_run.font.name = formatting_options['font_family']
            biblio_heading_run.font.bold = True
            
            # Phân tích văn bản và tìm các trích dẫn tiềm năng
            doc_nlp = nlp(content)
            potential_citations = []
            
            # Tìm các mẫu trích dẫn phù hợp với kiểu đã chọn
            citation_pattern = None
            if formatting_options['citation_style'] == 'apa':
                citation_pattern = re.compile(r'\(([^)]+), (\d{4})[^)]*\)')
            elif formatting_options['citation_style'] == 'mla':
                citation_pattern = re.compile(r'\(([^)]+) (\d+)[^)]*\)')
            elif formatting_options['citation_style'] == 'chicago':
                citation_pattern = re.compile(r'\[\d+\]')
            elif formatting_options['citation_style'] == 'ieee':
                citation_pattern = re.compile(r'\[(\d+)\]')
            
            if citation_pattern:
                for match in citation_pattern.finditer(content):
                    potential_citations.append(match.group(0))
            
            # Thêm các trích dẫn tìm được vào phần tài liệu tham khảo
            added_citations = set()
            for cite in potential_citations:
                if cite not in added_citations:
                    biblio_entry = doc.add_paragraph()
                    biblio_entry.add_run(f"{cite} - Reference details")
                    added_citations.add(cite)

        doc.save(output_path)
        print(f"Đã lưu tài liệu thành công vào: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"Error in document creation: {str(e)}")
        return None 